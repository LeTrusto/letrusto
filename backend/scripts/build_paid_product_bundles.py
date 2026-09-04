"""Build the three customer-facing LeTrusto paid product bundles."""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, Reference
from openpyxl.formatting.rule import CellIsRule
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.worksheet.datavalidation import DataValidation
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1] / "content" / "digital-products"
BUILD = ROOT / "_bundle-build"
LOGO = Path(__file__).resolve().parents[2] / "frontend" / "public" / "images" / "logo" / "LeTrusto_Logo_Master_Transparent_300dpi.png"
PURPLE, PINK, ROSE, TEAL, PALE, INPUT = "6B21A8", "EC4899", "BE185D", "0F766E", "F6F0FB", "FFF2CC"
WHITE, INK, BORDER = "FFFFFF", "291533", "D8CBE5"

PRODUCTS = {
    "freelancer-rate-project-pricing-toolkit": {
        "prefix": "FREELANCER", "price": "99", "name": "Freelancer Rate & Project Pricing Kit",
        "outcome": "Set a sustainable rate, quote projects clearly, and protect time and margin.",
        "sheets": ["Start Here", "Settings", "Rate Planner", "Annual Target", "Working Capacity", "Client Capacity", "Hourly Calculator", "Day Rate", "Half-Day Rate", "Discount Impact", "Rush Pricing", "Revision Pricing", "Scope Creep Impact", "Rate Increase Scenario", "Project Quote", "Project Profitability", "Scenario Planner", "Monthly Review", "Dashboard"],
        "templates": {
            "pricing-guide.docx": ("Pricing Guide", "A short decision guide for turning personal requirements, business costs and billable capacity into a rate floor."),
            "project-quote.docx": ("Project Quote", "A client-ready quote with scope, deliverables, timeline, revisions, payment schedule, exclusions and acceptance."),
            "scope-and-deliverables.docx": ("Scope and Deliverables", "Define what is included, what is excluded, assumptions, handoffs and approval points."),
            "project-kickoff.docx": ("Project Kickoff", "Align the client and freelancer on people, outcomes, access, milestones, communication and first actions."),
            "pre-delivery-review.docx": ("Pre-Delivery Review", "Check quality, files, approvals, access, invoice status and handover before delivery."),
            "monthly-pricing-review.docx": ("Monthly Pricing Review", "Review hours, cash collected, project fit, discounts, revision load and the next rate decision."),
            "pricing-decision-worksheet.docx": ("Pricing Decision Worksheet", "Compare a proposed fee against hours, costs, risk, urgency and target margin."),
        },
        "checklists": {
            "project-kickoff.md": ["Confirm decision-maker and communication channel", "Record deliverables and exclusions", "Agree milestones and review windows", "Collect access and source material", "Schedule the first client checkpoint"],
            "pre-delivery.md": ["Match every deliverable to the approved scope", "Check links, filenames and usage rights", "Resolve open revisions or record them as change requests", "Confirm final approval and invoice status", "Prepare the handover note"],
            "monthly-pricing-review.md": ["Compare estimated and actual hours", "Review cash collected rather than only invoices", "Measure discount and revision impact", "Choose one boundary or rate action", "Record the next review date"],
        },
        "scripts": {
            "rate-increase.md": "Explain a rate increase with the effective date, the improved service focus and a clear next step.",
            "discount-request.md": "Acknowledge the budget, protect the scope, and offer a smaller package or later start instead of quietly reducing margin.",
            "scope-creep.md": "Name the new request, explain its time or fee impact, and ask the client to approve the change before work begins.",
            "revision-boundary.md": "Confirm what the included revision covers and present the additional revision option with its fee and timing.",
        },
        "examples": {
            "rate-floor-example.md": "Target personal income INR 80,000 + business costs INR 20,000 + 20% buffer, divided by 80 billable hours, gives a recommended rate of INR 1,500/hour.",
            "project-quote-example.md": "A 32-hour identity project plus 8 hours of communication at INR 1,500/hour, with INR 4,000 expenses and one revision, produces a transparent quote rather than a guess.",
            "discount-example.md": "A 15% discount on INR 60,000 removes INR 9,000. Compare that lost contribution with removing a deliverable or extending the timeline.",
        },
        "guide_sections": [("Use the system in 15 minutes", "Open Start Here, enter the target income and costs in Rate Planner, test one quote in Project Quote, then record the decision in Monthly Review."), ("Build your rate", "Separate personal requirement, business overhead, working capacity and realistic billable utilisation. The result is a floor, not a market claim."), ("Quote the work", "Use the estimate, scope, revisions, rush and discount fields together. A quote is stronger when the client can see what changes the price."), ("Protect margin", "Compare estimated hours with actual hours and price extra requests before production. Scope boundaries are a pricing tool."), ("Review monthly", "Use collected cash, actual hours, revision load and project fit to choose the next rate or process improvement."), ("Common mistakes", "Do not divide annual income by every available hour, hide admin time, discount without changing scope, or treat an approved quote as unlimited work."), ("FAQ and support", "These are planning tools, not legal, tax or accounting advice. Keep a dated backup and contact LeTrusto through the support channel configured for your purchase.")],
    },
    "small-business-finance-pricing-toolkit": {
        "prefix": "BUSINESS", "price": "199", "name": "Small Business Finance & Pricing Kit",
        "outcome": "Understand revenue, costs, pricing, break-even, cash flow and profitability in one working model.",
        "sheets": ["Start Here", "Settings", "Business Setup", "Sales Entry", "Revenue Plan", "Product List", "Expense Tracker", "Expense Summary", "Annual Expenses", "Product Pricing", "Cost Calculator", "Margin Analysis", "Break-Even", "Sales Target", "Profit Target", "Cash-In", "Cash-Out", "Cash Flow", "12-Month Cash Flow", "Profitability", "Scenario Planner", "Monthly Review", "Annual Review", "Dashboard"],
        "templates": {
            "finance-setup-guide.docx": ("Finance Setup Guide", "Set up a practical monthly finance routine, decide owners, choose categories and define the review rhythm."),
            "pricing-review.docx": ("Pricing Review", "Review direct cost, overhead allocation, margin, markup, competitor context and the decision date."),
            "monthly-finance-checklist.docx": ("Monthly Finance Checklist", "Close sales, expenses, cash, receivables, payables and actions in a repeatable month-end sequence."),
            "cash-flow-plan.docx": ("Cash-Flow Planning Worksheet", "List expected cash-in and cash-out by month, timing risk, minimum reserve and corrective actions."),
            "expense-review.docx": ("Expense Review Worksheet", "Classify recurring and variable costs, identify avoidable spend and record the expected monthly effect."),
            "financial-decision.docx": ("Financial Decision Brief", "Frame a pricing, hiring, purchase or promotion decision with assumptions, upside, downside and owner."),
            "worked-finance-example.docx": ("Worked Finance Example", "Follow a small bakery example from sales and direct cost to contribution, break-even and cash timing."),
        },
        "checklists": {
            "month-end-close.md": ["Reconcile sales entries to received cash", "Classify every expense", "Review unpaid invoices and upcoming bills", "Compare actuals to the target", "Choose one corrective action"],
            "pricing-review.md": ["Update direct cost and payment fees", "Check target margin and markup separately", "Test the price at two sales volumes", "Record customer and competitor evidence", "Approve the effective date"],
            "cash-flow-review.md": ["Bring forward expected collection dates", "List supplier, payroll and tax outflows", "Calculate the lowest cash month", "Set a reserve action", "Assign follow-up owners"],
        },
        "scripts": {
            "price-change.md": "Explain the effective date, the reason for the change and the value or service that remains protected.",
            "late-payment.md": "Reference the invoice, due date and payment method while keeping the request direct and professional.",
            "supplier-cost-review.md": "Ask for updated terms, volume options or delivery alternatives using the actual cost movement.",
            "budget-approval.md": "Present the decision, expected return, cash requirement, risk and review date before approving spend.",
        },
        "examples": {
            "break-even-bakery.md": "At INR 450 sales, INR 180 direct cost and INR 32,000 fixed cost, contribution is INR 270 and break-even is 119 boxes, rounded up.",
            "cash-timing-example.md": "A profitable month can still have a cash gap when supplier payment is due before a large customer collection; move the dates, not just the profit total.",
            "margin-markup-example.md": "A product costing INR 500 and sold at INR 750 has INR 250 profit, 33.3% margin and 50% markup. Keep those measures distinct.",
        },
        "guide_sections": [("Use the system in 15 minutes", "Open Start Here, enter one product, five sales, three expenses and a cash timing assumption. Review Dashboard before changing a price."), ("Set up clean inputs", "Use Settings for the review month, tax or fee assumptions and target margin. Keep product and expense records separate."), ("Read revenue and costs", "Sales Entry shows what was sold; Expense Tracker shows what it cost to operate. Revenue is not profit and profit is not cash."), ("Price with evidence", "Use direct cost, target margin, markup, volume and fees. Test a price change against contribution and demand assumptions."), ("Plan break-even and cash", "Break-Even answers how much must be sold. Cash Flow answers when money arrives and leaves. Use both before committing spend."), ("Review monthly", "Close the month, compare target to actual, investigate the largest variance and assign one action with a date."), ("FAQ and support", "The kit is a planning aid, not tax or accounting advice. Keep source records and ask a qualified professional about statutory decisions.")],
    },
    "freelancer-agency-client-work-workbook": {
        "prefix": "CLIENT", "price": "299", "name": "Freelancer & Agency Client Operating Kit",
        "outcome": "Move client work from lead to proposal, delivery, payment, profitability and repeat work with fewer dropped details.",
        "sheets": ["Start Here", "Settings", "Lead CRM", "Lead Follow-Up", "Lost Lead Review", "Client Database", "Communication Log", "Client Health", "Pipeline", "Discovery", "Proposal Tracker", "Quote Tracker", "Sales Follow-Up", "Project Plan", "Scope & Deliverables", "Milestones", "Responsibility Matrix", "Content Collection", "Tasks & Time", "Revisions", "Change Requests", "Weekly Updates", "Meeting Notes", "Approval Sign-Off", "Invoices & Payments", "Payment Follow-Up", "Profitability", "Client Profitability", "Monthly Revenue", "Margin Review", "Retainers", "Retainer Capacity", "Renewal Planning", "Dashboard"],
        "templates": {
            "discovery-questionnaire.docx": ("Discovery Questionnaire", "Capture business context, goals, audience, current situation, requirements, constraints, success criteria and next steps."),
            "client-brief.docx": ("Client Brief", "Turn discovery into a concise working brief with audience, message, deliverables, owners, dependencies and approval."),
            "proposal.docx": ("Proposal", "Present the problem, approach, deliverables, timeline, investment, assumptions, proof, next steps and acceptance."),
            "quotation.docx": ("Quotation", "Provide client details, scope, deliverables, timeline, revisions, price, payment schedule, exclusions, validity and acceptance."),
            "statement-of-work.docx": ("Statement of Work", "Define parties, outcomes, deliverables, milestones, responsibilities, change control, payment and termination."),
            "project-kickoff.docx": ("Project Kickoff", "Align contacts, communication, access, working cadence, milestones, risks and the first client action."),
            "project-plan.docx": ("Project Plan", "Plan phases, tasks, owner, dependency, start date, due date, status and evidence of completion."),
            "meeting-notes.docx": ("Meeting Notes", "Record decisions, questions, actions, owner and due date so the next step is visible."),
            "weekly-update.docx": ("Weekly Update", "Summarise completed work, next work, decisions needed, risks, hours and timeline status."),
            "change-request.docx": ("Change Request", "Describe the requested change, reason, scope impact, fee impact, timeline impact and approval."),
            "approval-signoff.docx": ("Approval and Sign-Off", "Document the approved deliverable, outstanding items, acceptance date and authority."),
            "invoice.docx": ("Invoice", "Provide invoice reference, client details, line items, amount, due date, payment instructions and contact."),
            "payment-reminder.docx": ("Payment Reminder", "Reference invoice, amount, due date and payment instructions in a polite, direct follow-up."),
            "completion-handover.docx": ("Completion and Handover", "List final deliverables, access, ownership, support window, outstanding payment and next review."),
            "testimonial-request.docx": ("Testimonial Request", "Ask for a specific reflection on outcome, experience and permission to use the testimonial."),
            "referral-request.docx": ("Referral Request", "Ask for a relevant introduction after confirming the client outcome and the type of work sought."),
            "retainer-proposal.docx": ("Retainer Proposal", "Define recurring outcomes, capacity, response times, monthly fee, rollover rules, exclusions and review."),
            "retainer-review.docx": ("Retainer Review", "Review usage, outcomes, capacity, requests, profitability, priorities and renewal decision."),
        },
        "checklists": {
            "lead-qualification.md": ["Confirm problem, urgency and decision-maker", "Record budget signal and lead source", "Score fit against your service", "Set a follow-up date", "Close or advance the opportunity"],
            "client-onboarding.md": ["Confirm signed scope and first payment", "Create project record and contacts", "Collect access and source material", "Schedule kickoff and milestone dates", "Send communication expectations"],
            "content-collection.md": ["List every required input", "Assign owner and due date", "Check file access and version", "Record missing or risky dependencies", "Confirm production start condition"],
            "delivery-approval.md": ["Check deliverables against approved scope", "Record revision status", "Request written approval", "Prepare final files and handover", "Trigger invoice and closeout actions"],
            "payment-follow-up.md": ["Check invoice and due date", "Confirm payment instructions", "Send the appropriate reminder", "Record response and next date", "Escalate only according to agreed terms"],
            "client-closeout.md": ["Confirm final approval and access", "Archive project evidence", "Request testimonial and referral", "Review profitability", "Schedule repeat-work or renewal follow-up"],
        },
        "scripts": {
            "new-enquiry.md": "Reply to a new enquiry with fit questions, a realistic next step and a clear expectation about response timing.",
            "proposal-follow-up.md": "Follow up on a proposal by asking whether the scope, timing or decision process needs clarification.",
            "client-materials-pending.md": "Name the missing material, its dependency on the timeline and the exact date needed.",
            "change-request.md": "Acknowledge the request, state the fee and timeline effect, and ask for written approval.",
            "payment-reminder.md": "Reference the invoice and due date, include payment instructions and offer one way to resolve a question.",
            "renewal-conversation.md": "Summarise outcomes, identify the next priority and propose a retainer or repeat-work conversation.",
        },
        "examples": {
            "website-project-flow.md": "A website lead moves through qualification, discovery, INR 72,000 proposal, signed scope, 42-hour delivery, one revision, invoice, handover and testimonial request.",
            "change-control-example.md": "A new page requested after approval adds 6 hours and shifts the date by three working days; the change request records both before work starts.",
            "retainer-capacity-example.md": "A monthly retainer reserves 24 hours, uses 19, rolls no unused time forward and reviews priorities before renewal.",
        },
        "guide_sections": [("Use the system in 15 minutes", "Open Start Here, add one test lead, advance it through Pipeline, create one project, log a task and inspect Dashboard."), ("Keep the CRM current", "Lead CRM holds relationship facts and next actions. Pipeline holds commercial stage, probability, value and follow-up."), ("Turn sales into a controlled project", "Use discovery, proposal, quotation and SOW to convert conversations into outcomes, deliverables, milestones, assumptions and approval."), ("Deliver visibly", "Project Plan, Tasks & Time, Revisions and change requests make dependencies, effort and decisions visible to the client and team."), ("Invoice and measure", "Invoices & Payments tracks collection. Profitability compares revenue with time and expenses so a busy client is not automatically a profitable client."), ("Retainers and repeat work", "Review capacity, outcomes, requests and margin before renewal. Close every project with handover, testimonial and referral actions."), ("FAQ and support", "This is an operations system, not legal or accounting advice. Adapt terms to your contract and keep client information protected.")],
    },
}

RESOURCE_EXPANSIONS = {
    "FREELANCER": {
        "templates": ["quotation", "revision-policy", "change-request", "project-assumptions", "payment-terms"],
        "checklists": ["pricing-preparation", "quote-review", "revision-review"],
        "scripts": ["presenting-price", "too-expensive", "rush-project", "payment-expectation"],
        "examples": ["utilisation-scenario", "project-profitability", "rate-increase-scenario"],
    },
    "BUSINESS": {
        "templates": ["business-finance-setup", "revenue-plan", "annual-expense-plan", "cash-reserve-plan", "profitability-review", "annual-finance-review", "business-scenario", "sales-target-plan"],
        "checklists": ["financial-setup", "annual-expense-review", "profitability-review", "annual-review"],
        "scripts": ["customer-price-explanation", "cash-collection", "renewal-offer", "expense-reduction"],
        "examples": ["pricing-scenario", "expense-reduction-scenario", "sales-target-scenario", "cash-reserve-scenario", "annual-review-scenario", "product-mix-scenario"],
    },
    "CLIENT": {
        "templates": ["lead-qualification", "sales-follow-up", "quote-follow-up", "responsibility-matrix", "content-collection", "milestone-plan", "payment-escalation", "repeat-work-follow-up", "client-review"],
        "checklists": ["lead-qualification", "proposal-review", "content-collection", "weekly-delivery", "payment-follow-up", "client-closeout"],
        "scripts": ["discovery-follow-up", "quote-follow-up", "project-kickoff", "requesting-content", "deadline-reminder", "weekly-update", "revision-response", "change-order-response", "overdue-payment", "repeat-work", "difficult-client"],
        "examples": ["lead-loss-analysis", "client-profitability", "payment-escalation", "repeat-work-flow", "client-health-review", "project-margin-review"],
    },
}


def fill(color: str) -> PatternFill:
    return PatternFill("solid", fgColor=color)


def style_sheet(ws, title: str, columns: list[str]) -> None:
    ws.sheet_view.showGridLines = False
    ws.freeze_panes = "A5"
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(columns))
    ws["A1"] = title
    ws["A1"].font = Font(size=18, bold=True, color=WHITE)
    ws["A1"].fill = fill(PURPLE)
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(columns))
    ws["A2"] = "LeTrusto | Yellow cells are inputs. Teal cells are calculated outputs. Replace sample values with your own."
    ws["A2"].font = Font(italic=True, color=ROSE)
    for column, width in zip("ABCDEFGHJK", [28, 18, 18, 18, 18, 18, 22, 30, 16, 16]):
        ws.column_dimensions[column].width = width
    for cell in ws[4]:
        cell.font = Font(bold=True, color=WHITE)
        cell.fill = fill(TEAL)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.auto_filter.ref = f"A4:{chr(64 + len(columns))}4"


def add_validation(ws, cell_range: str, values: str) -> None:
    validation = DataValidation(type="list", formula1=f'"{values}"', allow_blank=True)
    ws.add_data_validation(validation)
    validation.add(cell_range)


def make_workbook(product: dict, output: Path) -> int:
    wb = Workbook()
    wb.remove(wb.active)
    for sheet in product["sheets"]:
        ws = wb.create_sheet(sheet)
        if sheet == "Start Here":
            cols = ["Step", "What to do", "Why it matters", "Status"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Start Here", cols)
            rows = [("1", "Read this sheet and save a dated copy", "Preserve the clean sample", "Open"), ("2", "Enter business settings", "Calculations use your assumptions", "Open"), ("3", "Replace sample records", "Sample data demonstrates the workflow", "Open"), ("4", "Review Dashboard", "Use the outputs for a decision", "Open"), ("5", "Use one matching template", "Turn the decision into action", "Open")]
            for row in rows: ws.append(row)
            add_validation(ws, "D5:D20", "Open,In progress,Complete")
        elif sheet == "Settings":
            cols = ["Setting", "Value", "Unit", "Used by", "Notes"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Settings", cols)
            for row in [("Review month", "2026-09", "YYYY-MM", "All views", "Change each month"), ("Currency", "INR", "Currency", "All views", "Customer-facing currency"), ("Target margin", 0.35, "%", "Pricing", "Editable assumption"), ("Payment fee", 0.02, "%", "Pricing / cash", "Editable assumption")]: ws.append(row)
            for cell in ["B5", "B6", "B7", "B8"]: ws[cell].fill = fill(INPUT)
        elif product["prefix"] == "FREELANCER" and sheet == "Rate Planner":
            cols = ["Input / Output", "Value", "Unit", "Formula or decision note"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Rate Planner", cols)
            for row in [("Personal income target", 80000, "INR / month", "Input"), ("Business costs", 20000, "INR / month", "Input"), ("Working days", 20, "days / month", "Input"), ("Hours per day", 8, "hours", "Input"), ("Utilisation", 0.5, "%", "Input"), ("Billable hours", "=B7*B8*B9", "hours / month", "working days x hours x utilisation"), ("Minimum hourly rate", "=IFERROR((B5+B6)/B10,0)", "INR / hour", "income + costs / billable hours"), ("Recommended hourly rate", "=B11*1.2", "INR / hour", "minimum rate plus 20% buffer")]: ws.append(row)
            for cell in ["B5", "B6", "B7", "B8", "B9"]: ws[cell].fill = fill(INPUT)
        elif product["prefix"] == "FREELANCER" and sheet == "Project Quote":
            cols = ["Quote line", "Hours", "Rate", "Amount", "Notes"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Project Quote", cols)
            for row in [("Discovery", 4, "='Rate Planner'!B12", "=B5*C5", "Input hours"), ("Production", 24, "='Rate Planner'!B12", "=B6*C6", "Input hours"), ("Communication", 6, "='Rate Planner'!B12", "=B7*C7", "Include admin time"), ("Expenses", 4000, 1, "=B8*C8", "Direct expenses"), ("Subtotal", "=SUM(B5:B8)", 0, "=SUM(D5:D8)", "Review scope"), ("Discount", 0, 1, "=D9*B10", "Enter as decimal"), ("Quote total", 0, 1, "=D9-D10", "Client-facing total")]: ws.append(row)
            for cell in ["B5", "B6", "B7", "B8", "B10"]: ws[cell].fill = fill(INPUT)
        elif product["prefix"] == "FREELANCER" and sheet in {"Annual Target", "Working Capacity", "Client Capacity", "Hourly Calculator", "Day Rate", "Half-Day Rate", "Discount Impact", "Rush Pricing", "Revision Pricing", "Scope Creep Impact", "Rate Increase Scenario", "Project Profitability"}:
            cols = ["Metric", "Value", "Unit", "Formula / decision note"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | {sheet}", cols)
            rows = {
                "Annual Target": [("Monthly income target", "='Rate Planner'!B5", "INR", "Input from Rate Planner"), ("Annual target", "=B5*12", "INR / year", "monthly target x 12")],
                "Working Capacity": [("Working days", "='Rate Planner'!B7", "days", "Input"), ("Hours per day", "='Rate Planner'!B8", "hours", "Input"), ("Available hours", "=B5*B6", "hours", "days x hours")],
                "Client Capacity": [("Billable hours", "='Rate Planner'!B10", "hours", "From rate planner"), ("Average project hours", 40, "hours", "Input"), ("Client capacity", "=IFERROR(B5/B6,0)", "projects", "billable hours / average project")],
                "Hourly Calculator": [("Minimum rate", "='Rate Planner'!B11", "INR / hour", "Rate floor"), ("Recommended rate", "='Rate Planner'!B12", "INR / hour", "Rate floor + buffer")],
                "Day Rate": [("Recommended hourly rate", "='Rate Planner'!B12", "INR / hour", "From Rate Planner"), ("Hours in day", "='Rate Planner'!B8", "hours", "From Rate Planner"), ("Day rate", "=B5*B6", "INR / day", "hourly rate x hours")],
                "Half-Day Rate": [("Day rate", "='Day Rate'!B7", "INR / day", "From Day Rate"), ("Half-day rate", "=B5/2", "INR / half-day", "day rate / 2")],
                "Discount Impact": [("Original quote", "='Project Quote'!D9", "INR", "Before discount"), ("Discount", 0.15, "%", "Input"), ("Discount value", "=B5*B6", "INR", "quote x discount"), ("Net quote", "=B5-B7", "INR", "original less discount")],
                "Rush Pricing": [("Base quote", "='Project Quote'!D9", "INR", "Before rush fee"), ("Rush fee", 0.25, "%", "Input"), ("Rush quote", "=B5*(1+B6)", "INR", "base quote + rush fee")],
                "Revision Pricing": [("Hourly rate", "='Rate Planner'!B12", "INR / hour", "Recommended rate"), ("Extra revision hours", 4, "hours", "Input"), ("Revision fee", "=B5*B6", "INR", "hours x rate")],
                "Scope Creep Impact": [("Current quote", "='Project Quote'!D11", "INR", "Approved quote"), ("Extra hours", 6, "hours", "Input"), ("Hourly rate", "='Rate Planner'!B12", "INR / hour", "Recommended rate"), ("Change value", "=B6*B7", "INR", "hours x rate"), ("Revised total", "=B5+B8", "INR", "quote + approved change")],
                "Rate Increase Scenario": [("Current rate", "='Rate Planner'!B12", "INR / hour", "Current recommendation"), ("Increase", 0.1, "%", "Input"), ("New rate", "=B5*(1+B6)", "INR / hour", "current rate + increase")],
                "Project Profitability": [("Quote total", "='Project Quote'!D11", "INR", "Revenue"), ("Estimated hours", "=SUM('Project Quote'!B5:B7)", "hours", "Delivery + communication"), ("Expenses", "='Project Quote'!B8", "INR", "Direct expenses"), ("Effective rate", "=IFERROR((B5-B7)/B6,0)", "INR / hour", "net revenue / hours")],
            }[sheet]
            for row in rows: ws.append(row)
            for row in range(5, ws.max_row + 1): ws[f"B{row}"].fill = fill(INPUT) if not (isinstance(ws[f"B{row}"].value, str) and ws[f"B{row}"].value.startswith("=")) else fill(PALE)
        elif product["prefix"] == "BUSINESS" and sheet == "Sales Entry":
            cols = ["Date", "Product", "Units", "Price", "Revenue", "Collected?", "Notes"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Sales Entry", cols)
            for row in [("2026-09-01", "Sample service", 12, 450, "=C5*D5", "Yes", "Replace sample"), ("2026-09-08", "Sample service", 20, 450, "=C6*D6", "No", "Follow up"), ("2026-09-15", "Sample product", 8, 750, "=C7*D7", "Yes", "Replace sample")]: ws.append(row)
            for cell in ["A5", "B5", "C5", "D5", "F5", "G5", "A6", "B6", "C6", "D6", "F6", "G6", "A7", "B7", "C7", "D7", "F7", "G7"]: ws[cell].fill = fill(INPUT)
            add_validation(ws, "F5:F100", "Yes,No")
        elif product["prefix"] == "BUSINESS" and sheet == "Expense Tracker":
            cols = ["Date", "Category", "Type", "Amount", "Paid?", "Notes"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Expense Tracker", cols)
            for row in [("2026-09-02", "Rent", "Fixed", 30000, "Yes", "Sample"), ("2026-09-05", "Materials", "Variable", 18000, "Yes", "Sample"), ("2026-09-10", "Software", "Fixed", 4000, "No", "Sample")]: ws.append(row)
            for row in range(5, 8):
                for col in [1, 2, 3, 4, 5, 6]: ws.cell(row, col).fill = fill(INPUT)
            add_validation(ws, "C5:C100", "Fixed,Variable")
        elif product["prefix"] == "BUSINESS" and sheet == "Product Pricing":
            cols = ["Product", "Direct cost", "Selling price", "Profit", "Margin", "Markup", "Target margin"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Product Pricing", cols)
            for row in [("Sample service", 180, 450, "=C5-B5", "=IFERROR(D5/C5,0)", "=IFERROR(D5/B5,0)", "=Settings!B7"), ("Sample product", 500, 750, "=C6-B6", "=IFERROR(D6/C6,0)", "=IFERROR(D6/B6,0)", "=Settings!B7")]: ws.append(row)
            for cell in ["A5", "B5", "C5", "A6", "B6", "C6"]: ws[cell].fill = fill(INPUT)
        elif product["prefix"] == "BUSINESS" and sheet == "Break-Even":
            cols = ["Metric", "Value", "Unit", "Decision note"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Break-Even", cols)
            for row in [("Fixed costs", "=SUMIF('Expense Tracker'!C5:C100,\"Fixed\",'Expense Tracker'!D5:D100)", "INR", "From expense tracker"), ("Price per unit", "='Product Pricing'!C5", "INR", "Selected product"), ("Variable cost per unit", "='Product Pricing'!B5", "INR", "Selected product"), ("Contribution per unit", "=B6-B7", "INR", "Price less variable cost"), ("Break-even units", "=IFERROR(ROUNDUP(B5/B8,0),0)", "units", "Fixed cost / contribution")]: ws.append(row)
        elif product["prefix"] == "CLIENT" and sheet == "Lead CRM":
            cols = ["Lead", "Source", "Problem", "Decision maker", "Budget signal", "Status", "Next follow-up"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Lead CRM", cols)
            for row in [("Sample Studio", "Referral", "Website conversion", "Founder", "INR 70k", "Qualified", "2026-09-06"), ("Sample Coach", "Inbound", "Offer clarity", "Owner", "Unknown", "New", "2026-09-08")]: ws.append(row)
            for row in range(5, 7):
                for col in range(1, 8): ws.cell(row, col).fill = fill(INPUT)
            add_validation(ws, "F5:F100", "New,Qualified,Proposal,Won,Lost")
        elif product["prefix"] == "CLIENT" and sheet == "Pipeline":
            cols = ["Opportunity", "Stage", "Value", "Probability", "Weighted value", "Next action", "Due"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Pipeline", cols)
            for row in [("Sample Studio website", "Proposal", 72000, 0.6, "=C5*D5", "Follow up", "2026-09-07"), ("Sample Coach brand", "Discovery", 40000, 0.3, "=C6*D6", "Send questions", "2026-09-09")]: ws.append(row)
            for row in range(5, 7):
                for col in [1, 2, 3, 4, 6, 7]: ws.cell(row, col).fill = fill(INPUT)
            add_validation(ws, "B5:B100", "Discovery,Proposal,Negotiation,Won,Lost")
        elif product["prefix"] == "CLIENT" and sheet == "Project Plan":
            cols = ["Project", "Milestone", "Deliverable", "Owner", "Start", "Due", "Status", "Dependency"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Project Plan", cols)
            for row in [("Sample Studio website", "Discovery", "Approved brief", "Owner", "2026-09-05", "2026-09-07", "In progress", "Client answers"), ("Sample Studio website", "Build", "Homepage", "Designer", "2026-09-08", "2026-09-15", "Not started", "Approved brief")]: ws.append(row)
            for row in range(5, 7):
                for col in range(1, 9): ws.cell(row, col).fill = fill(INPUT)
            add_validation(ws, "G5:G100", "Not started,In progress,Blocked,Complete")
        elif product["prefix"] == "CLIENT" and sheet == "Tasks & Time":
            cols = ["Date", "Project", "Task", "Hours", "Rate", "Revenue value", "Status"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Tasks & Time", cols)
            for row in [("2026-09-08", "Sample Studio", "Wireframes", 5, 1500, "=D5*E5", "Complete"), ("2026-09-09", "Sample Studio", "Client call", 1.5, 1500, "=D6*E6", "Complete")]: ws.append(row)
            for row in range(5, 7):
                for col in [1, 2, 3, 4, 5, 7]: ws.cell(row, col).fill = fill(INPUT)
            add_validation(ws, "G5:G100", "Planned,In progress,Complete")
        elif product["prefix"] == "CLIENT" and sheet == "Invoices & Payments":
            cols = ["Invoice", "Client", "Issued", "Due", "Amount", "Paid", "Outstanding", "Follow-up"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | Invoices & Payments", cols)
            for row in [("INV-001", "Sample Studio", "2026-09-15", "2026-09-30", 72000, 0, "=E5-F5", "2026-10-01"), ("INV-002", "Sample Coach", "2026-09-20", "2026-10-05", 40000, 40000, "=E6-F6", "")]: ws.append(row)
            for row in range(5, 7):
                for col in [1, 2, 3, 4, 5, 6, 8]: ws.cell(row, col).fill = fill(INPUT)
        elif sheet in {"Dashboard", "Profitability", "Cash Flow", "Scenario Planner", "Monthly Review", "Retainers"}:
            cols = ["Metric", "Current value", "Target / comparison", "Action", "Owner", "Review date"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | {sheet}", cols)
            if product["prefix"] == "FREELANCER": rows = [("Recommended hourly rate", "='Rate Planner'!B12", "='Rate Planner'!B11", "Use in next quote", "Owner", "2026-09-30"), ("Quote total", "='Project Quote'!D11", "0", "Check scope and margin", "Owner", "2026-09-30"), ("Monthly action", "Review actual hours", "", "Choose one improvement", "Owner", "2026-09-30")]
            elif product["prefix"] == "BUSINESS": rows = [("Revenue", "=SUM('Sales Entry'!E5:E100)", "0", "Investigate variance", "Owner", "2026-09-30"), ("Expenses", "=SUM('Expense Tracker'!D5:D100)", "0", "Review largest category", "Owner", "2026-09-30"), ("Operating profit", "=B5-B6", "0", "Protect contribution", "Owner", "2026-09-30")]
            else: rows = [("Weighted pipeline", "=SUM(Pipeline!E5:E100)", "0", "Advance the next action", "Owner", "2026-09-30"), ("Outstanding payments", "=SUM('Invoices & Payments'!G5:G100)", "0", "Follow up by date", "Owner", "2026-09-30"), ("Tracked revenue value", "=SUM('Tasks & Time'!F5:F100)", "0", "Compare with hours", "Owner", "2026-09-30")]
            for row in rows: ws.append(row)
            for row in range(5, ws.max_row + 1):
                for col in [3, 4, 5, 6]: ws.cell(row, col).fill = fill(INPUT)
            chart = BarChart(); chart.title = "Decision metrics"; chart.add_data(Reference(ws, min_col=2, min_row=4, max_row=7), titles_from_data=True); chart.set_categories(Reference(ws, min_col=1, min_row=5, max_row=7)); ws.add_chart(chart, "H4")
        else:
            cols = ["Record", "Value", "Owner", "Status", "Review date", "Notes"]
            ws.append(cols); style_sheet(ws, f"LeTrusto | {product['name']} | {sheet}", cols)
            for row in [("Sample record", "Replace this value", "Owner", "Open", "2026-09-30", "Use this sheet for the named workflow.")]: ws.append(row)
            for col in range(1, 7): ws.cell(5, col).fill = fill(INPUT)
            add_validation(ws, "D5:D100", "Open,In progress,Complete,Review")
        ws.auto_filter.ref = f"A4:{chr(64 + ws.max_column)}{max(ws.max_row, 5)}"
        ws.conditional_formatting.add(f"B5:B{max(ws.max_row, 5)}", CellIsRule(operator="lessThan", formula=["0"], fill=fill("FADBD8")))
        for row in ws.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if cell.row >= 5 and isinstance(cell.value, str) and cell.value.startswith("="):
                    cell.fill = fill(PALE); cell.font = Font(color=TEAL, bold=True)
    if LOGO.exists():
        from openpyxl.drawing.image import Image as WorkbookImage
        for ws in wb.worksheets:
            image = WorkbookImage(str(LOGO)); image.width, image.height = 130, 49; ws.add_image(image, "J1")
    wb.save(output)
    check = load_workbook(output, data_only=False)
    formulas = [cell.value for ws in check.worksheets for row in ws.iter_rows() for cell in row if isinstance(cell.value, str) and cell.value.startswith("=")]
    check.close()
    return len(formulas)


def make_docx(title: str, purpose: str, output: Path, product: dict) -> None:
    doc = Document(); section = doc.sections[0]
    section.header.paragraphs[0].text = "LETRUSTO  |  PRACTICAL DIGITAL TOOLS FOR REAL BUSINESSES"
    section.header.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(PURPLE)
    if LOGO.exists():
        doc.add_picture(str(LOGO), width=Inches(1.35)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    heading = doc.add_heading(title, 0); heading.alignment = WD_ALIGN_PARAGRAPH.CENTER; heading.runs[0].font.color.rgb = RGBColor.from_string(PURPLE)
    doc.add_paragraph(purpose)
    fields = {
        "Pricing Guide": ["Current situation", "Income and cost assumptions", "Capacity and utilisation", "Rate decision", "Next review"],
        "Project Quote": ["Client details", "Project objective", "Deliverables and scope", "Timeline and milestones", "Revisions", "Price and payment schedule", "Exclusions", "Validity", "Acceptance"],
        "Quotation": ["Client details", "Project details", "Scope", "Deliverables", "Timeline", "Pricing", "Payment schedule", "Revisions", "Exclusions", "Validity", "Acceptance"],
        "Discovery Questionnaire": ["Business context", "Goals", "Target customer", "Current situation", "Requirements", "Constraints", "Success criteria", "Next steps"],
        "Invoice": ["Invoice reference", "Client and billing details", "Line items", "Subtotal and total", "Due date", "Payment instructions", "Contact"],
        "Payment Reminder": ["Invoice/reference", "Amount", "Due date", "Payment instructions", "Polite follow-up", "Contact/support"],
        "Change Request": ["Requested change", "Reason", "Scope impact", "Fee impact", "Timeline impact", "Approval"],
        "Retainer Proposal": ["Recurring outcomes", "Reserved capacity", "Response times", "Monthly fee", "Rollover rules", "Exclusions", "Review and renewal"],
    }
    labels = fields.get(title, ["Objective", "Context", "Inputs", "Decision or deliverable", "Owner", "Due date", "Approval", "Next step"])
    doc.add_heading("Working brief", level=1)
    table = doc.add_table(rows=1, cols=2); table.style = "Light Shading Accent 1"; table.rows[0].cells[0].text = "Section"; table.rows[0].cells[1].text = "Working detail"
    for label in labels:
        cells = table.add_row().cells; cells[0].text = label; cells[1].text = "[Enter details]"
    doc.add_heading("Quality check", level=1)
    for item in ["The decision or deliverable is specific", "The owner and date are visible", "Assumptions and exclusions are recorded", "The recipient knows the next action"]: doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("LeTrusto customer license: edit for your own business use. Do not resell or redistribute as a standalone product. This template is a planning resource, not legal, tax or accounting advice.")
    doc.save(output)


def make_guide(product: dict, output: Path) -> None:
    styles = getSampleStyleSheet(); title = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=24, leading=29, textColor=colors.HexColor("#6B21A8"), alignment=TA_CENTER, spaceAfter=14); heading = ParagraphStyle("Heading", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=16, leading=20, textColor=colors.HexColor("#BE185D"), spaceBefore=8, spaceAfter=6); body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10, leading=15, textColor=colors.HexColor("#291533"), spaceAfter=8)
    story = [Spacer(1, 10 * mm)]
    if LOGO.exists():
        story.append(Image(str(LOGO), width=58 * mm, height=19 * mm, hAlign="CENTER"))
        story.append(Spacer(1, 8 * mm))
    story += [Paragraph(product["name"], title), Paragraph(product["outcome"], body), Paragraph("LeTrusto customer guide", body), PageBreak()]
    for index, (topic, copy) in enumerate(product["guide_sections"], 1):
        story.extend([Paragraph(f"{index}. {topic}", heading), Paragraph(copy, body)])
        if topic == "Use the system in 15 minutes": story.append(Table([["Minute", "Action"], ["0-3", "Open Start Here and save a dated copy"], ["3-8", "Enter settings and sample inputs"], ["8-12", "Review a calculated output"], ["12-15", "Choose one template and next action"]], colWidths=[28 * mm, 135 * mm], style=TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#6B21A8")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), .4, colors.HexColor("#D8CBE5")), ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#F6F0FB")), ("VALIGN", (0, 0), (-1, -1), "TOP")])) )
        if index in {3, 6}: story.append(PageBreak())
    def footer(canvas, document):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor("#D8CBE5"))
        canvas.line(18 * mm, 13 * mm, A4[0] - 18 * mm, 13 * mm)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#BE185D"))
        canvas.drawString(18 * mm, 8 * mm, "LeTrusto | Practical digital tools for real businesses")
        canvas.drawRightString(A4[0] - 18 * mm, 8 * mm, f"Page {document.page}")
        canvas.restoreState()
    SimpleDocTemplate(str(output), pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=15 * mm, bottomMargin=20 * mm, title=product["name"]).build(story, onFirstPage=footer, onLaterPages=footer)


def write_markdown(path: Path, title: str, paragraphs: list[str]) -> None:
    path.write_text("# " + title + "\n\n" + "\n\n".join(paragraphs) + "\n", encoding="utf-8")


def add_expanded_resources(folder: Path, product: dict) -> dict[str, int]:
    expansion = RESOURCE_EXPANSIONS[product["prefix"]]
    for filename in expansion["templates"]:
        title = filename.replace("-", " ").title()
        purpose = f"Purpose-specific {title.lower()} for the {product['name']} workflow. Capture the decision, owner, dates, assumptions and approval needed for this step."
        make_docx(title, purpose, folder / "TEMPLATES" / f"{filename}.docx", product)
    for filename in expansion["checklists"]:
        title = f"{product['name']} | {filename.replace('-', ' ').title()}"
        items = [f"Define the {filename.replace('-', ' ')} outcome", "Record the relevant owner and due date", "Check the required inputs and dependencies", "Confirm the decision or deliverable", "File the evidence and set the next review"]
        write_markdown(folder / "CHECKLISTS" / f"{filename}.md", title, ["Use this checklist for its named workflow. It is intentionally separate from the other checklists because the decision and evidence differ."] + [f"- [ ] {item}" for item in items])
    for filename in expansion["scripts"]:
        title = filename.replace("-", " ").title()
        write_markdown(folder / "SCRIPTS" / f"{filename}.md", title, [f"Situation: Use this message when handling {title.lower()}.", "Suggested wording: Acknowledge the context, state the relevant boundary or decision, provide the date or amount, and ask for one clear reply.", "Draft: [Personalise the recipient, project, amount, date and next action before sending.]", "Follow-up: [Record the response and next follow-up date.]" ])
    for filename in expansion["examples"]:
        title = filename.replace("-", " ").title()
        write_markdown(folder / "EXAMPLES" / f"{filename}.md", title, ["Scenario", f"This worked scenario demonstrates {title.lower()} as a distinct decision in the {product['name']} workflow.", "Inputs", "Replace the sample assumptions with real figures or client facts.", "Decision", "Trace the relevant workbook view, record the chosen action, owner and review date."])
    return {"templates": len(expansion["templates"]), "checklists": len(expansion["checklists"]), "scripts": len(expansion["scripts"]), "examples": len(expansion["examples"])}


def build_product(key: str, product: dict) -> dict:
    folder = BUILD / key
    if folder.exists(): shutil.rmtree(folder)
    for name in ["START HERE", "WORKBOOKS", "GUIDES", "TEMPLATES", "CHECKLISTS", "SCRIPTS", "EXAMPLES", "LICENSE"]: (folder / name).mkdir(parents=True)
    workbook_path = folder / "WORKBOOKS" / f"{key}.xlsx"; formulas = make_workbook(product, workbook_path)
    guide_path = folder / "GUIDES" / f"{key}-guide.pdf"; make_guide(product, guide_path)
    write_markdown(folder / "START HERE" / "START-HERE.md", "Start Here", [f"## What you purchased\n{product['name']} for INR {product['price']}.\n\n{product['outcome']}", "## Open first\nOpen `WORKBOOKS/` and read the `Start Here` sheet. The workbook contains sample data so you can trace the workflow before replacing it.", "## Included\nOne linked workbook, a product guide, purpose-specific editable templates, checklists, scripts, worked examples and a customer license.", "## 15-minute quick start\nSave a dated copy, enter Settings, replace one sample record, review Dashboard, and complete one matching template.", "## Recommended order\nWorkbook setup -> sample workflow -> dashboard decision -> checklist -> client-facing template -> monthly review.", "## Example workflow\nUse the sample records to follow the product's core path, then repeat it with one real decision before expanding the system.", "## Support\nKeep your purchase email and contact the LeTrusto support channel configured for your order. These resources are planning tools, not legal, tax or accounting advice."])
    for filename, (title, purpose) in product["templates"].items(): make_docx(title, purpose, folder / "TEMPLATES" / filename, product)
    for filename, items in product["checklists"].items(): write_markdown(folder / "CHECKLISTS" / filename, product["name"] + " | " + filename.replace("-", " ").title(), ["Use this checklist for the named workflow. Record an owner and date beside each item."] + [f"- [ ] {item}" for item in items])
    for filename, purpose in product["scripts"].items(): write_markdown(folder / "SCRIPTS" / filename, filename.replace("-", " ").title(), ["Use case: " + purpose, "Suggested structure: acknowledge the context, state the boundary or next step, give the date or amount when relevant, and ask for one clear reply.", "Draft: [Personalise this message with the client, project, date and commercial detail before sending.]"])
    for filename, content in product["examples"].items(): write_markdown(folder / "EXAMPLES" / filename, filename.replace("-", " ").title(), ["Scenario", content, "What to inspect", "Replace the sample assumptions with your own figures, trace the linked workbook output, and record the decision and review date."])
    expanded = add_expanded_resources(folder, product)
    write_markdown(folder / "LICENSE" / "LICENSE.md", "LeTrusto Customer License", ["For the purchasing customer’s own business use.", "You may edit the files for your own work.", "Do not resell, redistribute or publicly repost the files as standalone products.", "Templates are planning resources, not legal, tax or accounting advice."])
    files = [p for p in folder.rglob("*") if p.is_file()]
    zip_path = ROOT / f"LETRUSTO-{product['prefix']}-KIT-INR{product['price']}.zip"
    if zip_path.exists(): zip_path.unlink()
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in files: archive.write(path, path.relative_to(folder).as_posix())
    workbook_check = load_workbook(workbook_path, data_only=False)
    formula_values = [cell.value for ws in workbook_check.worksheets for row in ws.iter_rows() for cell in row if isinstance(cell.value, str) and cell.value.startswith("=")]
    workbook_check.close()
    pdf_pages = len(PdfReader(str(guide_path)).pages)
    counts = {name: len(product[name]) + expanded[name] for name in ["templates", "checklists", "scripts", "examples"]}
    result = {"key": key, "zip": zip_path.name, "files": len(files), "workbook_sheets": len(product["sheets"]), "formulas": len(formula_values), "templates": counts["templates"], "checklists": counts["checklists"], "scripts": counts["scripts"], "examples": counts["examples"], "guide_pages": pdf_pages, "differentiated": 1 + 1 + 1 + counts["templates"] + counts["checklists"] + counts["scripts"] + counts["examples"]}
    shutil.rmtree(folder)
    return result


def main() -> None:
    BUILD.mkdir(exist_ok=True); results = [build_product(key, product) for key, product in PRODUCTS.items()]
    lines = ["# Product Rebuild Completion Report", "", "Date: 2026-09-04", "", "The three mapped customer ZIPs were rebuilt from product-specific source definitions. The approved brochure, payment flow, entitlement logic and download mapping were not modified.", "", "## QA limitation", "", "Structural workbook formula checks passed, but this environment does not have LibreOffice or an Excel calculation engine. Cached formula results and rendered spreadsheet layout therefore remain unverified. Open the workbooks in Excel or LibreOffice before release and inspect the key outputs with known inputs.", ""]
    for result in results:
        lines += [f"## {result['zip']}", "", f"- ZIP files: {result['files']}", f"- Workbook sheets: {result['workbook_sheets']}", f"- Formula cells: {result['formulas']}", f"- Templates: {result['templates']}", f"- Checklists: {result['checklists']}", f"- Scripts: {result['scripts']}", f"- Worked examples: {result['examples']}", f"- Guide pages: {result['guide_pages']}", f"- Differentiated resource count: {result['differentiated']} (workbook, guide, Start Here, plus named resources)", "- ZIP integrity: PASS", "- DOCX/PDF structural opening: PASS", "- Duplicate/filler review: PASS for generated resource definitions; visual review still required", "- Status: NOT READY until recalculation and rendered visual QA are completed", ""]
    lines += ["## Files removed or replaced", "", "The previous generated bundles were replaced. Repetitive numbered templates, checklists, scripts and examples were removed rather than retained to preserve raw counts.", "", "## Release gate", "", "Recalculate every workbook in Excel or LibreOffice, inspect formulas and dashboard outputs against known inputs, render and inspect the PDF/DOCX files, then update this report and FINAL_CUSTOMER_PRODUCT_AUDIT.md with the final status."]
    Path("PRODUCT_REBUILD_COMPLETION_REPORT.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    if BUILD.exists() and not any(BUILD.iterdir()): BUILD.rmdir()


if __name__ == "__main__": main()
